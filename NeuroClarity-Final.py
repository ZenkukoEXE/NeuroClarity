import functools
import streamlit as st
import io
import re
import math
try:
    import PyPDF2
except ModuleNotFoundError as e:
    raise ModuleNotFoundError(
        "PyPDF2 is required to parse PDF uploads. Install it with `pip install PyPDF2` or add it to requirements.txt."
    ) from e
from docx import Document
import nltk
from nltk.corpus import wordnet
from nltk.tokenize import word_tokenize
import os
import warnings

# Suppress Streamlit warnings
warnings.filterwarnings("ignore", category=UserWarning, module="streamlit")

# --- 1. SETUP & NLTK DOWNLOADS ---
@functools.lru_cache(maxsize=1)
def setup_nltk():
    nltk.download('punkt')
    nltk.download('wordnet')
    nltk.download('averaged_perceptron_tagger')

setup_nltk()

# --- 2. THESAURUS SIMPLIFIER ---
def get_simpler_synonym(word):
    synonyms = []
    for syn in wordnet.synsets(word):
        for lemma in syn.lemmas():
            synonyms.append(lemma.name())
    
    if not synonyms:
        return word
        
    synonyms = list(set(synonyms))
    synonyms.sort(key=len) 
    
    # Return shortest synonym if shorter than original
    if len(synonyms[0]) < len(word):
        return synonyms[0].replace('_', ' ')
    return word

def simplify_text(text):
    words = word_tokenize(text)
    simplified_words = []
    for word in words:
        # Simplify longer words, keep punctuation intact
        if len(word) > 7 and word.isalpha():
            simplified_words.append(get_simpler_synonym(word))
        else:
            simplified_words.append(word)
            
    result = ' '.join(simplified_words)
    result = re.sub(r'\s([?.!,"\'](?:\s|$))', r'\1', result)
    return result

# --- 3. TIME ESTIMATION ---
def estimate_time(text):
    word_count = len(text.split())
    # Base: 10 mins per 50 words
    base_minutes = (word_count / 50) * 10 
    
    words = re.findall(r"\b\w+\b", text.lower())
    if 'task' not in words:
        return 0.0
        
    # Only the word 'task' affects hour calculation now
    base_minutes *= 2.0
    
    # Convert minutes to hours
    base_hours = base_minutes / 60
    return round(base_hours, 2) if base_hours > 0 else 0.0

# --- 4. FILE EXTRACTION ---
def extract_text(uploaded_file):
    text = ""
    if uploaded_file.name.endswith('.pdf'):
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n\n"
    elif uploaded_file.name.endswith('.docx'):
        doc = Document(uploaded_file)
        for para in doc.paragraphs:
            text += para.text + "\n\n"
    elif uploaded_file.name.endswith('.txt'):
        text = str(uploaded_file.read(), "utf-8")
    return text

# --- 5. DOCUMENT GENERATOR ---
def create_downloadable_docx(processed_sections):
    doc = Document()
    doc.add_heading('Simplified Academic Brief', 0)
    
    for section in processed_sections:
        doc.add_heading(section['Title'], level=1)
        doc.add_paragraph(f"⏱️ Suggested Time: {section['Time']:.2f} hours").bold = True
        doc.add_heading('Simplified Instructions:', level=2)
        doc.add_paragraph(section['Simplified'])
        doc.add_page_break()
        
    # Save to a virtual file (bytes) so the user can download it directly
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# --- 6. USER INTERFACE (STREAMLIT) ---
st.title("🧠 NeuroClarity: Academic Brief Simplifier")
st.write("Upload your academic brief (PDF, DOCX, or TXT). We will simplify the complex language and provide a time-management plan!")

uploaded_file = st.file_uploader("Upload Document", type=['pdf', 'docx', 'txt'])

if uploaded_file is not None:
    with st.spinner('Extracting and analyzing text...'):
        raw_text = extract_text(uploaded_file)
        
        # Simple sectioning: split by double line breaks (paragraphs)
        raw_sections = [sec.strip() for sec in raw_text.split('\n\n') if len(sec.strip()) > 20]
        
        processed_sections = []
        for i, section_text in enumerate(raw_sections):
            processed_sections.append({
                "Title": f"Section {i+1}",
                "Original": section_text,
                "Simplified": simplify_text(section_text),
                "Time": estimate_time(section_text)
            })

    st.success("Brief processed successfully!")

    # Add user input for total assignment time
    st.sidebar.header("Customize Time Allocation")
    user_total_time = st.sidebar.number_input(
        "Enter total time available for assignment (in hours):",
        min_value=0.1, step=0.1, value=1.0
    )  # Default to 1 hour

    # Adjust time allocation based on user input
    total_time = sum([s['Time'] for s in processed_sections])

    if user_total_time > 0:
        scaling_factor = user_total_time / total_time
        for sec in processed_sections:
            sec['Time'] *= scaling_factor

    st.info(f"Total Estimated Time to Complete Assignment: {user_total_time:.2f} hours")

    with st.expander("View Section Breakdown"):
        for sec in processed_sections:
            st.subheader(sec['Title'])
            st.write(f"**Adjusted Time:** {sec['Time']:.2f} hours")
            st.write(f"**Simplified Text:** {sec['Simplified']}")
            st.divider()

    # Create the downloadable file
    docx_file = create_downloadable_docx(processed_sections)
    
    st.download_button(
        label="📥 Download Simplified Brief (.docx)",
        data=docx_file,
        file_name="Simplified_Brief.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

